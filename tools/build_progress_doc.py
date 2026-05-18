from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Echo项目构建与阶段进度说明.docx"


ACCENT = "4F46E5"
ACCENT_LIGHT = "EEF2FF"
TEXT = "111827"
MUTED = "4B5563"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margin(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = tc_pr.first_child_found_in("w:tcMar")
    if margin is None:
        margin = OxmlElement("w:tcMar")
        tc_pr.append(margin)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margin.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margin.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string(ACCENT if level <= 2 else TEXT)
    return p


def add_body(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(TEXT)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(TEXT)


def add_placeholder(doc, number, title, guidance):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT_LIGHT)
    set_cell_border(cell, ACCENT, "10")
    set_cell_margin(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"截图位置 {number}：{title}")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(guidance)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], ACCENT)
        set_cell_border(hdr[i], ACCENT)
        set_cell_margin(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in hdr[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_border(cells[i])
            set_cell_margin(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor.from_string(TEXT)
        set_cell_shading(cells[0], "F8FAFC")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_timeline_item(doc, stage, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(stage)
    r1.bold = True
    r1.font.name = "Microsoft YaHei"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r1.font.color.rgb = RGBColor.from_string(ACCENT)
    r1.font.size = Pt(10.5)
    r2 = p.add_run(f"  {title}")
    r2.bold = True
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(10.5)
    add_body(doc, body)


def configure_styles(doc):
    styles = doc.styles
    for style_name in ("Normal", "Body Text", "List Bullet"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    configure_styles(doc)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Echo 项目构建与阶段进度说明")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    r = subtitle.add_run("从工程搭建到当前完成度的过程记录（含截图占位）")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ("项目名称", "Echo 轻量级音频分享社区"),
        ("技术栈", "Django / SQLite / Tailwind CSS / HTMX / 原生 JavaScript"),
        ("文档日期", "2026-05-16"),
        ("当前验证", "Django check 通过；35 个自动化测试全部通过"),
    ]
    for row, item in zip(meta.rows, items):
        row.cells[0].text = item[0]
        row.cells[1].text = item[1]
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9.5)
        set_cell_shading(row.cells[0], ACCENT_LIGHT)
    doc.add_paragraph()

    add_placeholder(
        doc,
        1,
        "项目首页或整体应用壳",
        "建议截取浏览器首页完整首屏，展示顶部栏、左侧音乐库、中部内容、右侧播放上下文和底部播放器。",
    )

    doc.add_page_break()

    add_heading(doc, "1. 项目定位与建设目标", 1)
    add_body(
        doc,
        "Echo 是一个基于 Django 的轻量级音频分享社区课程项目。项目从静态界面原型逐步演进为可运行的本地 Web 应用，目标是在有限复杂度内完成音频上传、浏览、播放、歌词、评论、用户资料和基础社区互动的闭环。",
    )
    add_bullets(
        doc,
        [
            "产品侧目标：提供类似音乐社区的浏览、上传、播放、评论和歌词体验。",
            "工程侧目标：建立清晰的 Django app 边界、可维护的模板结构、稳定的全局播放器状态和基础回归测试。",
            "演示侧目标：支持本地启动、演示数据填充、常见页面 200 检查和自动化测试验证。",
        ],
    )

    add_heading(doc, "2. 项目构建过程", 1)
    add_timeline_item(
        doc,
        "阶段 1",
        "工程骨架搭建",
        "创建 Django 项目与基础 app，拆分 settings 为 base/dev/prod，建立 templates、static、media、路由和本地 SQLite 开发环境。项目启动方式统一为 manage.py runserver，并保留 seed_demo 命令用于演示数据初始化。",
    )
    add_timeline_item(
        doc,
        "阶段 2",
        "核心数据模型接入",
        "围绕 Track、Album、TrackComment、TrackLyrics、TrackLyricLine 等实体建立 ORM 模型和迁移，首页、列表页、详情页、评论页和歌词页逐步从静态数据切换到数据库查询。",
    )
    add_timeline_item(
        doc,
        "阶段 3",
        "应用壳与播放器成型",
        "实现 Spotify 风格三栏应用壳：顶部导航、左侧音乐库、中部主内容、右侧播放上下文、底部常驻播放器。播放器支持播放/暂停、上一首/下一首、进度、音量、循环、随机和本地状态持久化。",
    )
    add_timeline_item(
        doc,
        "阶段 4",
        "上传、歌词与评论闭环",
        "音频上传支持常见格式、封面主题、基础元数据读取和安全校验；歌词支持 LRC/TXT 上传与解析；评论支持按曲目读取、排序筛选和 HTMX partial 提交。",
    )
    add_timeline_item(
        doc,
        "阶段 5",
        "用户系统与资料能力",
        "接入 Django 注册、登录、退出、保持登录和用户资料页。头像上传加入前端裁剪预览、后端文件头校验和压缩处理，数据库仅保存文件路径。",
    )
    add_timeline_item(
        doc,
        "阶段 6",
        "交互细节与稳定性修复",
        "持续修复播放器队列、右键菜单、重复歌曲入队、右侧播放列表自动展开、中文编码、布局折叠按钮和页面缓存等问题，同时补充基础回归测试。",
    )

    add_placeholder(
        doc,
        2,
        "项目目录结构",
        "建议截取 IDE 左侧目录树，重点展示 tracks、albums、comments、lyrics、core、templates、static、docs、tests 等目录。",
    )

    add_heading(doc, "3. 当前系统结构", 1)
    add_table(
        doc,
        ["模块", "职责", "当前状态"],
        [
            ("core", "首页、应用壳上下文、评论/歌词主页面聚合、SQLite 初始化", "已落地"),
            ("tracks", "音频作品、上传、列表、详情、播放量接口、上传安全校验", "已落地"),
            ("albums", "专辑/合辑容器、专辑与曲目关系", "基础能力已落地"),
            ("comments", "评论模型、评论查询、评论提交、评论反应模型", "主体已落地，点赞与回复 UI 待完善"),
            ("lyrics", "歌词版本、歌词行、LRC/TXT 解析、歌词上传", "主体已落地，仍有细节修复项"),
            ("auth/profile", "注册、登录、退出、用户资料、头像处理", "主体已落地"),
            ("tests", "页面、上传、安全、播放、评论、歌词等回归测试", "35 个测试通过"),
        ],
        widths=[3.0, 8.2, 4.0],
    )

    add_heading(doc, "4. 已完成功能概览", 1)
    add_table(
        doc,
        ["功能域", "已完成内容", "说明"],
        [
            ("音频浏览与播放", "首页推荐、最新列表、详情页、播放量上报、全局播放器", "支持顺序、随机、列表循环、单曲循环、上一首/下一首"),
            ("播放队列", "真实 playQueue 状态、队列下标、重复歌曲入队、右侧队列展示", "同一首歌可在播放列表中重复出现"),
            ("右键菜单", "按场景区分列表、搜索、详情页、播放队列等菜单", "下一首播放、加入队列、移除队列等动作已接入前端状态"),
            ("歌词", "歌词上传、解析、同步高亮、点击歌词 seek", "支持 available / instrumental / pending 状态"),
            ("评论", "按曲目读取评论、筛选排序、HTMX 提交、禁用缓存", "回复与点赞模型已有基础，UI 仍需继续完成"),
            ("用户", "注册、登录、退出、保持登录、资料页、头像上传裁剪", "上传内容会绑定当前用户"),
            ("工程质量", "settings 拆分、错误页、空状态、分页、上传校验、回归测试", "当前 manage.py check 与 35 个测试通过"),
        ],
        widths=[3.2, 7.0, 5.0],
    )

    add_placeholder(
        doc,
        3,
        "音频上传页面",
        "建议截取 /tracks/upload/ 页面，展示音频文件、封面、歌词录入或上传校验提示。",
    )
    add_placeholder(
        doc,
        4,
        "播放器与右侧播放列表",
        "建议截取播放歌曲后的右侧播放列表，尤其展示重复歌曲入队、当前播放高亮和队列数量。",
    )

    add_heading(doc, "5. 数据库与安全策略", 1)
    add_body(
        doc,
        "现阶段数据库采用 Django ORM + SQLite。为了适配播放量、评论、点赞等写入场景，项目启动时对 SQLite 连接启用 WAL、NORMAL synchronous 和 busy_timeout。计数字段更新采用 F() 表达式，避免并发写入覆盖。",
    )
    add_bullets(
        doc,
        [
            "Track 默认仅展示 published 状态，隐藏、草稿和删除状态不会进入普通页面读取。",
            "上传文件同时检查扩展名、文件大小、文件头魔数和空文件，避免只靠后缀判断。",
            "头像上传限制 JPG / PNG / WEBP，后端继续做图片解析和压缩。",
            "评论页禁用缓存，切歌或筛选时重新按 track 参数读取数据库。",
        ],
    )

    add_heading(doc, "6. 交互与前端应用壳", 1)
    add_body(
        doc,
        "前端采用 Tailwind CSS、HTMX 和原生 JavaScript。整体不是完整 SPA，但通过常驻播放器、#main-content 局部替换和本地状态持久化，形成接近桌面音乐应用的使用体验。",
    )
    add_bullets(
        doc,
        [
            "左侧音乐库支持折叠，创建菜单包含上传音乐、歌单、共享合辑、文件夹和上传歌词入口。",
            "右侧上下文支持正在播放与播放列表视图，播放列表变化时可主动展开队列视图。",
            "底部播放器负责跨页面状态，歌词和评论按钮会跟随当前播放曲目更新 track 参数。",
            "顶部、左侧、右侧和底部控件的布局仍在持续打磨，最近重点修复了中文编码和折叠按钮占位问题。",
        ],
    )
    add_placeholder(
        doc,
        5,
        "歌词页同步播放效果",
        "建议播放一首带 LRC 的歌曲后截取歌词页，展示当前行高亮、背景色和底部播放器。",
    )
    add_placeholder(
        doc,
        6,
        "评论页交互",
        "建议截取 /comments/?track=<id> 页面，展示热门/最新/问题筛选、评论表单和评论列表。",
    )

    add_heading(doc, "7. 当前完成度评估", 1)
    add_table(
        doc,
        ["状态", "内容", "评估"],
        [
            ("已完成", "工程骨架、主要 app、数据库模型、上传、播放、歌词、评论、认证、头像、基础错误页和测试", "具备课程项目演示能力"),
            ("基本完成", "全局播放器、播放队列、右键菜单、右侧上下文、响应式应用壳", "核心体验可用，仍需继续做浏览器级回归"),
            ("进行中", "搜索结果体验、评论回复 UI、评论点赞、播放队列视图完全由 playQueue 驱动", "已有入口或部分基础，仍待补全"),
            ("待建设", "歌单 CRUD、通知系统、管理后台、完整个人主页动态、前端自动化测试", "下一阶段重点"),
            ("风险点", "base.html 与 echo-shell.js 体量较大，中文编码和复杂布局容易回归", "修改前后需要做针对性页面验证"),
        ],
        widths=[2.6, 8.2, 4.4],
    )

    add_heading(doc, "8. 验证记录", 1)
    add_body(doc, "本次文档生成前，对当前工作区执行了基础验证，结果如下。")
    add_table(
        doc,
        ["验证项", "命令", "结果"],
        [
            ("Django 系统检查", r".\.venv\Scripts\python.exe manage.py check", "通过，未发现 issues"),
            ("自动化测试", r".\.venv\Scripts\python.exe manage.py test tests", "35 个测试全部通过"),
            ("JS 语法检查", r"node --check static\js\echo-shell.js", "近期多次通过；当前未发现语法错误"),
            ("首页访问", "http://127.0.0.1:8000/", "本地服务可返回 200"),
        ],
        widths=[3.0, 8.0, 4.0],
    )
    add_placeholder(
        doc,
        7,
        "测试结果或终端验证",
        "建议截取 manage.py check 和 manage.py test tests 的终端输出，作为阶段验收证据。",
    )

    add_heading(doc, "9. 下一阶段建议", 1)
    add_bullets(
        doc,
        [
            "优先修复审查待跟进项：歌词 source_file 保存前 seek(0)、评论回复命名冲突、重复脚本抽取。",
            "补齐搜索系统：结果分页、按类型筛选、搜索结果播放队列和空状态。",
            "完善评论互动：回复 UI、点赞接口、TrackCommentReaction 前后端闭环。",
            "实现歌单 CRUD：创建、编辑、排序、添加/移除曲目，并区分歌单与临时播放队列。",
            "补充浏览器级自动化测试，覆盖播放、切歌、歌词、评论、右键菜单和响应式布局。",
            "继续拆分 base.html 和 echo-shell.js，降低后续功能修改的回归风险。",
        ],
    )
    add_placeholder(
        doc,
        8,
        "后续计划标注图",
        "建议截取 issue/TODO 文档或项目看板，用于展示下一阶段任务安排。",
    )

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    add_page_number(section.footer.paragraphs[0])
    add_heading(doc, "附录：常用命令", 1)
    add_table(
        doc,
        ["用途", "命令"],
        [
            ("启动本地服务", r".\.venv\Scripts\python.exe manage.py runserver 127.0.0.1:8000"),
            ("访问地址", "http://127.0.0.1:8000/"),
            ("填充演示数据", r".\.venv\Scripts\python.exe manage.py seed_demo"),
            ("Django 检查", r".\.venv\Scripts\python.exe manage.py check"),
            ("自动化测试", r".\.venv\Scripts\python.exe manage.py test tests"),
            ("JS 语法检查", r"node --check static\js\echo-shell.js"),
        ],
        widths=[4.0, 11.0],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
